from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, List


def _chapter_title(chapter: Dict[str, Any]) -> str:
    title = str(chapter.get("title") or "").strip()
    if title:
        return title
    idx = int(chapter.get("index") or 0)
    return f"Chapter {idx}" if idx else "Chapter"


def _iter_blocks(chunk: Dict[str, Any]) -> List[Dict[str, Any]]:
    blocks = chunk.get("translated_blocks")
    if isinstance(blocks, list) and blocks:
        return [b for b in blocks if isinstance(b, dict)]
    text = str(chunk.get("translation") or "").strip()
    if not text:
        return []
    return [{"type": "paragraph", "translated_text": text, "footnote_number": ""}]


def _collect_render_units(translated_chunks: Iterable[Dict[str, Any]]) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    body: List[Dict[str, Any]] = []
    footnotes: List[Dict[str, Any]] = []
    chunks = sorted(
        list(translated_chunks),
        key=lambda x: (int(x.get("chapter_index") or 0), int(x.get("chunk_index") or 0)),
    )
    for chunk in chunks:
        for block in _iter_blocks(chunk):
            typ = str(block.get("type") or "paragraph")
            text = str(block.get("translated_text") or "").strip()
            if not text:
                continue
            rec = {
                "chapter_index": int(chunk.get("chapter_index") or 0),
                "chunk_index": int(chunk.get("chunk_index") or 0),
                "type": typ,
                "text": text,
                "footnote_number": str(block.get("footnote_number") or "").strip(),
            }
            if typ == "footnote":
                footnotes.append(rec)
            else:
                body.append(rec)
    return body, footnotes


def render_translation_text(state: Dict[str, Any], translated_chunks: Iterable[Dict[str, Any]]) -> str:
    body, footnotes = _collect_render_units(translated_chunks)
    by_chapter: Dict[int, List[Dict[str, Any]]] = {}
    for block in body:
        by_chapter.setdefault(int(block.get("chapter_index") or 0), []).append(block)

    chapters = state.get("chapters") or []
    title = str(state.get("source_name") or state.get("project_id") or "translation").strip()
    target = str(state.get("target_language") or "").strip()

    parts: List[str] = [title]
    if target:
        parts.append(f"Target language: {target}")
    parts.append("")

    for chapter in chapters:
        chapter_index = int(chapter.get("index") or 0)
        chapter_chunks = by_chapter.get(chapter_index) or []
        if not chapter_chunks:
            continue
        parts.append(f"# {_chapter_title(chapter)}")
        parts.append("")
        for block in chapter_chunks:
            text = str(block.get("text") or "").strip()
            typ = str(block.get("type") or "paragraph")
            if not text:
                continue
            if typ == "heading":
                parts.append(f"## {text}")
            elif typ == "list_item":
                parts.append(text if text.lstrip().startswith(("-", "*", "•")) else f"- {text}")
            else:
                parts.append(text)
            parts.append("")

    if footnotes:
        parts.append("# 脚注")
        parts.append("")
        for idx, note in enumerate(footnotes, start=1):
            original = str(note.get("footnote_number") or idx)
            parts.append(f"[{idx}] {note.get('text')} (原注 {original})")
        parts.append("")

    return "\n".join(parts).strip() + "\n"


def export_txt(state: Dict[str, Any], translated_chunks: Iterable[Dict[str, Any]], output_path: str | Path) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(render_translation_text(state, translated_chunks), encoding="utf-8")
    return path


def export_docx(state: Dict[str, Any], translated_chunks: Iterable[Dict[str, Any]], output_path: str | Path) -> Path:
    try:
        from docx import Document  # type: ignore
    except ModuleNotFoundError as e:
        raise ModuleNotFoundError("Missing dependency for .docx export. Install with: pip install python-docx") from e

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    body, footnotes = _collect_render_units(translated_chunks)
    by_chapter: Dict[int, List[Dict[str, Any]]] = {}
    for block in body:
        by_chapter.setdefault(int(block.get("chapter_index") or 0), []).append(block)

    doc = Document()
    title = str(state.get("source_name") or state.get("project_id") or "Translation").strip()
    doc.add_heading(title, level=0)
    target = str(state.get("target_language") or "").strip()
    if target:
        doc.add_paragraph(f"Target language: {target}")

    for chapter in state.get("chapters") or []:
        chapter_index = int(chapter.get("index") or 0)
        chapter_chunks = by_chapter.get(chapter_index) or []
        if not chapter_chunks:
            continue
        doc.add_heading(_chapter_title(chapter), level=1)
        for block in chapter_chunks:
            text = str(block.get("text") or "").strip()
            if not text:
                continue
            typ = str(block.get("type") or "paragraph")
            if typ == "heading":
                doc.add_heading(text, level=2)
            elif typ == "list_item":
                doc.add_paragraph(text.lstrip("-*• ").strip(), style="List Bullet")
            else:
                doc.add_paragraph(text)

    if footnotes:
        doc.add_heading("脚注", level=1)
        for idx, note in enumerate(footnotes, start=1):
            original = str(note.get("footnote_number") or idx)
            doc.add_paragraph(f"[{idx}] {note.get('text')} (原注 {original})")

    doc.save(str(path))
    return path


def export_translation(
    state: Dict[str, Any],
    translated_chunks: Iterable[Dict[str, Any]],
    output_path: str | Path,
    *,
    output_format: str = "txt",
) -> Path:
    fmt = (output_format or "txt").strip().lower().lstrip(".")
    if fmt == "docx":
        return export_docx(state, translated_chunks, output_path)
    if fmt == "txt":
        return export_txt(state, translated_chunks, output_path)
    raise ValueError(f"Unsupported export format: {output_format!r} (expected txt or docx)")
